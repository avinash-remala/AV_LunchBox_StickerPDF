#!/usr/bin/env python3
"""
Update Word template with data extracted from image.
Replaces existing cell contents in the template table.

Usag                 # Pick the longest match (real names are longer than garbage like 'vy' or 'v')
            if all_name_matches:
                name = max(all_name_matches, key=len)
            else:
                name = ''
            
            # Remove leading single characters followed by space (residual garbage like "v Abhishek")
            if name and ' ' in name:
                parts = name.split()
                if len(parts[0]) == 1 and len(parts) > 1:
                    # Skip the single-letter garbage
                    name = ' '.join(parts[1:])Check if it's a box type
        box_match = box_pattern.search(line)
        if box_match:
            current_row['box_type'] = box_match.group(1)
            print(f"→ Found box type: {current_row['box_type']}")
            continueython update_template.py <input_image> <template_docx> <output_docx>

Example:
    python update_template.py input.png template.docx output.docx
"""

import sys
import re
from pathlib import Path
from datetime import datetime

# OCR and image processing
try:
    from PIL import Image
    import pytesseract
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", 
                          "pillow", "pytesseract", "--break-system-packages"])
    from PIL import Image
    import pytesseract

# Word document handling
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.dml import MSO_THEME_COLOR
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", 
                          "python-docx", "--break-system-packages"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.dml import MSO_THEME_COLOR


def extract_table_data_from_image(image_path):
    """
    Extract table data from image using OCR and image analysis.
    Returns list of dictionaries with keys: name, address, box_type, rice_type
    
    Strategy: The image has two sections:
    1. Orders section (name + address) - appears first
    2. Box/Rice section (box types + rice types) - appears later
    
    We extract each section separately, then match by index/position.
    """
    print(f"Loading image: {image_path}")
    img = Image.open(image_path)
    
    # Perform OCR to get full text
    print("Performing OCR on image...")
    text = pytesseract.image_to_string(img)
    
    # Split into lines
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    
    print(f"OCR extracted {len(lines)} lines total")
    if lines:
        print("First 10 lines for debugging:")
        for i, line in enumerate(lines[:10]):
            print(f"  {i}: {line}")
    
    # Patterns
    phone_pattern = re.compile(r'\b\d{10}\b')
    address_pattern = re.compile(r'(2900 Plano Pkwy|3400 W Plano Pkwy)')
    box_pattern = re.compile(r'((?:Veg|Non-Veg)\s+Comfort Box|(?:\d+\s+)?(?:Comfort Box|Kabuli Chana Box|Moong Dal Box|Rajma Box))')
    rice_pattern = re.compile(r'(Pulav Rice|White Rice)')
    
    # FIRST PASS: Extract orders (name + address)
    orders = []
    for idx, line in enumerate(lines):
        # Skip header and junk
        if any(skip in line for skip in ['OAN', 'RWDN', 'ORDER NO', 'DELIVERY', 'PHONE']):
            continue
        
        # Skip lines with too many repeated characters
        if len(line) > 0:
            repeated_chars = sum(1 for i in range(len(line)-1) if line[i] == line[i+1])
            if repeated_chars > len(line) / 3:
                continue
        
        # Check if line contains address
        address_match = address_pattern.search(line)
        if address_match:
            address = address_match.group(1)
            line_remainder = line.replace(address, '').strip()
            
            # Extract phone number first
            phone_match = phone_pattern.search(line_remainder)
            if phone_match:
                line_remainder = line_remainder.replace(phone_match.group(), '').strip()
            
            # Extract NAME from what's left
            cleaned = line_remainder
            
            # Remove leading garbage
            while cleaned and not cleaned[0].isalpha():
                cleaned = cleaned[1:].strip()
            
            # Remove garbage characters
            garbage_patterns = ['»', '|', '¥', '«', '{', '}']
            for pattern in garbage_patterns:
                cleaned = cleaned.replace(pattern, ' ')
            
            cleaned = cleaned.lstrip('_- ~').rstrip('_- ~{')
            cleaned = ' '.join(cleaned.split())
            
            # Extract name sequences
            import re as regex_module
            all_name_matches = regex_module.findall(r'[A-Za-z]+(?:\s+[A-Za-z]+)*', cleaned)
            
            if all_name_matches:
                name = max(all_name_matches, key=len)
            else:
                name = ''
            
            # Clean up common OCR artifacts in names
            # Remove leading garbage patterns like "vy", "v", "a", etc. (single/double letter junk)
            if name:
                parts = name.split()
                # Remove leading single or double-letter garbage words
                while parts and len(parts[0]) <= 2 and parts[0].lower() in ['v', 'vy', 'a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'w', 'x', 'y', 'z', 'aa', 'ab', 'ac', 'ad', 'ae', 'af', 'ag', 'ah', 'ai', 'aj', 'ak', 'al', 'am', 'an', 'ao', 'ap', 'aq', 'ar', 'as', 'at', 'au', 'av', 'aw', 'ax', 'ay', 'az']:
                    parts.pop(0)
                name = ' '.join(parts)
            
            # Handle concatenated names (like "SaiCharanKurella" -> "Sai Charan Kurella")
            # Split on capital letters if the name has no spaces and looks like multiple words
            if name and ' ' not in name and len(name) > 4:
                # Check if it looks like concatenated words (multiple capital letters)
                capital_positions = [i for i, c in enumerate(name) if c.isupper()]
                if len(capital_positions) >= 2:
                    # Split the name at capital letter positions
                    split_name = []
                    for i, pos in enumerate(capital_positions):
                        start = pos
                        end = capital_positions[i + 1] if i + 1 < len(capital_positions) else len(name)
                        split_name.append(name[start:end])
                    name = ' '.join(split_name)
            
            # Remove leading single letters followed by space (residual garbage)
            if name and ' ' in name:
                parts = name.split()
                if len(parts[0]) == 1 and len(parts) > 1:
                    name = ' '.join(parts[1:])
            
            # Only add if we have a real name
            if name and len(name) > 3:
                orders.append({'name': name, 'address': address})
    
    # SECOND PASS: Extract all box types and rice types (they appear later in the document)
    box_types = []
    rice_types = []
    for line in lines:
        box_match = box_pattern.search(line)
        if box_match:
            box_type = box_match.group(1)
            box_types.append(box_type)
        
        rice_match = rice_pattern.search(line)
        if rice_match:
            rice_type = rice_match.group(1)
            rice_types.append(rice_type)
    
    print(f"\nExtracted: {len(orders)} orders, {len(box_types)} box types, {len(rice_types)} rice types")
    
    # COMBINE: Match by index position
    # The assumption is that box_types[i] and rice_types[i] correspond to orders[i]
    num_rows = max(len(orders), len(box_types), len(rice_types))
    data_rows = []
    
    for i in range(num_rows):
        row_data = {
            'name': orders[i]['name'] if i < len(orders) else '',
            'address': orders[i]['address'] if i < len(orders) else '',
            'box_type': box_types[i] if i < len(box_types) else '',
            'rice_type': rice_types[i] if i < len(rice_types) else ''
        }
        data_rows.append(row_data)
        print(f"Row {i+1}: {row_data}")
    
    print(f"\nExtracted {len(data_rows)} complete rows from image")
    return data_rows


def set_cell_margins(cell, left=None, right=None, top=None, bottom=None):
    """
    Set cell margins (padding) in twentieths of a point (twips).
    1 point = 20 twips
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    margins = {
        'left': left,
        'right': right,
        'top': top,
        'bottom': bottom
    }
    
    for margin_name, value in margins.items():
        if value is not None:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    
    tcPr.append(tcMar)


def convert_docx_to_pdf(docx_path, pdf_path):
    """
    Convert DOCX to PDF using available tools on Mac/Linux/Windows
    """
    import subprocess
    import platform
    import shutil
    
    system = platform.system()
    
    # Method 1: Try LibreOffice (works on Mac, Linux, Windows)
    libreoffice_paths = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # Mac
        'soffice',  # Linux/Windows (if in PATH)
        '/usr/bin/soffice',  # Linux
        'C:\\Program Files\\LibreOffice\\program\\soffice.exe',  # Windows
    ]
    
    for soffice_path in libreoffice_paths:
        if shutil.which(soffice_path) or Path(soffice_path).exists():
            try:
                print(f"Converting with LibreOffice: {soffice_path}")
                output_dir = str(Path(pdf_path).parent)
                
                # Run LibreOffice conversion
                result = subprocess.run([
                    soffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    docx_path
                ], capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0:
                    # LibreOffice creates output.pdf in the output directory
                    # We need to rename it to our desired name
                    generated_pdf = str(Path(output_dir) / Path(docx_path).with_suffix('.pdf').name)
                    if Path(generated_pdf).exists() and generated_pdf != pdf_path:
                        import shutil
                        shutil.move(generated_pdf, pdf_path)
                    
                    if Path(pdf_path).exists():
                        print("PDF created successfully with LibreOffice")
                        return
                    
            except Exception as e:
                print(f"LibreOffice conversion failed: {e}")
                continue
    
    # Method 2: Try docx2pdf (Python package - works well on Windows)
    try:
        import docx2pdf
        print("Converting with docx2pdf...")
        docx2pdf.convert(docx_path, pdf_path)
        if Path(pdf_path).exists():
            print("PDF created successfully with docx2pdf")
            return
    except ImportError:
        print("docx2pdf not available (run: pip install docx2pdf)")
    except Exception as e:
        print(f"docx2pdf conversion failed: {e}")
    
    # Method 3: Try unoconv (Linux/Mac alternative)
    if shutil.which('unoconv'):
        try:
            print("Converting with unoconv...")
            subprocess.run([
                'unoconv',
                '-f', 'pdf',
                '-o', pdf_path,
                docx_path
            ], check=True)
            if Path(pdf_path).exists():
                print("PDF created successfully with unoconv")
                return
        except Exception as e:
            print(f"unoconv conversion failed: {e}")
    
    # If all methods fail, provide instructions
    print("\n" + "=" * 60)
    print("ERROR: Could not convert to PDF automatically")
    print("=" * 60)
    print("\nThe DOCX file was created successfully, but PDF conversion failed.")
    print("Please install one of the following:\n")
    
    if system == "Darwin":  # Mac
        print("Option 1 (Recommended): LibreOffice")
        print("  brew install libreoffice")
        print("\nOption 2: Use Preview")
        print("  Open the DOCX in Preview and export as PDF")
        print("\nOption 3: Use Pages")
        print("  Open the DOCX in Pages and export as PDF")
    elif system == "Linux":
        print("Option 1 (Recommended): LibreOffice")
        print("  sudo apt-get install libreoffice")
        print("\nOption 2: unoconv")
        print("  sudo apt-get install unoconv")
    elif system == "Windows":
        print("Option 1: docx2pdf")
        print("  pip install docx2pdf")
        print("\nOption 2: LibreOffice")
        print("  Download from: https://www.libreoffice.org/download/")
    
    print("\nThen run the script again.")
    print("=" * 60)
    
    # Keep the DOCX file for manual conversion
    final_docx = pdf_path.replace('.pdf', '.docx')
    if Path(docx_path).exists():
        import shutil
        shutil.copy(docx_path, final_docx)
        print(f"\nDOCX file saved as: {final_docx}")
        print("You can manually convert this to PDF.")


def get_marker_for_box_rice(box_type, rice_type):
    """
    Determine the marker and font size increase based on box and rice type combinations.
    
    Returns: (marker_string, font_size_increase_pt)
    - Veg Comfort Box + Pulav Rice: --- VP --- + 2pt
    - Non-Veg Comfort Box + Pulav Rice: --- NVP --- + 2pt
    - Veg Comfort Box + White Rice: --- VW --- + 2pt
    - Non-Veg Comfort Box + White Rice: --- NVW --- + 2pt
    """
    # Check for Non-Veg FIRST to avoid matching "Veg" in "Non-Veg"
    is_non_veg = "Non-Veg" in box_type
    is_veg = "Veg" in box_type and not is_non_veg
    is_comfort_box = "Comfort Box" in box_type
    is_pulav = "Pulav Rice" in rice_type
    
    if is_comfort_box:
        if is_veg and is_pulav:
            return "--- VP ---", 2
        elif is_non_veg and is_pulav:
            return "--- NVP ---", 2
        elif is_veg and not is_pulav:
            return "--- VW ---", 2
        elif is_non_veg and not is_pulav:
            return "--- NVW ---", 2
    
    return "", 0


def update_template_with_data(template_path, output_path, data_rows):
    """
    Update the Word template by replacing cell contents with extracted data.
    Handles pagination by adding new rows/pages as needed.
    Then convert to PDF.
    """
    from docx import Document
    from docx.oxml import parse_xml
    from copy import deepcopy
    
    print(f"Loading template: {template_path}")
    doc = Document(template_path)
    table = doc.tables[0]
    data_columns = [0, 2, 4]
    offset_2px = Pt(2)
    data_index = 0
    
    # Get the number of available rows in the template
    initial_row_count = len(table.rows)
    print(f"Template has {initial_row_count} rows initially")
    
    # Calculate how many additional rows we need
    rows_needed = len(data_rows)
    if rows_needed > initial_row_count:
        additional_rows_needed = rows_needed - initial_row_count
        print(f"Need to add {additional_rows_needed} new rows for {len(data_rows)} data items")
        
        # Get a reference row to copy (usually the first data row)
        if initial_row_count > 0:
            reference_row = table.rows[0]
            tbl = table._element
            
            # Add new rows by copying the reference row
            for _ in range(additional_rows_needed):
                # Deep copy the reference row
                new_row = deepcopy(reference_row._element)
                tbl.append(new_row)
                print(f"Added new row (total now: {len(table.rows)})")
    
    # Now fill all rows with data
    for row_idx, row in enumerate(table.rows):
        for col_idx in data_columns:
            if data_index < len(data_rows):
                data = data_rows[data_index]
                
                # Change + to - in the display text
                cell_text = f"{data['name']}\n{data['address']}\n{data['box_type']} - {data['rice_type']}"
                
                cell = row.cells[col_idx]
                cell.text = ""  # Clear existing text
                
                # Add the three lines of text
                for idx, line in enumerate(cell_text.split('\n')):
                    if idx > 0:
                        cell.add_paragraph()
                    para = cell.paragraphs[idx]
                    
                    # For the third line (box_type - rice_type), add markers with special formatting
                    if idx == 2:
                        # Clear the paragraph
                        para.clear()
                        
                        # Add the text
                        run = para.add_run(line)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Get marker based on box/rice combination
                        marker, font_size_increase = get_marker_for_box_rice(data['box_type'], data['rice_type'])
                        
                        # Get current font size (default to 11pt if not set)
                        current_font_size = run.font.size
                        if current_font_size is None:
                            current_font_size = Pt(11)
                        else:
                            # Convert to Pt if it's in twips (EMU)
                            if isinstance(current_font_size, int):
                                current_font_size = Pt(current_font_size / 100)
                        
                        # Add marker in a separate paragraph if applicable
                        if marker:
                            # Create a new paragraph for markers
                            marker_para = cell.add_paragraph()
                            marker_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            marker_run = marker_para.add_run(marker)
                            marker_run.bold = True
                            # Set marker font size to base + 2pt
                            marker_run.font.size = current_font_size + Pt(font_size_increase)
                            
                            print(f"  Line 3: Added marker '{marker}' in separate paragraph (bold, +{font_size_increase}pt) for {data['box_type']} - {data['rice_type']}")
                    elif idx == 0:
                        # Name line - make it bold
                        para.clear()
                        run = para.add_run(line)
                        run.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        current_indent = para.paragraph_format.left_indent or Pt(0)
                        if col_idx == 2:
                            para.paragraph_format.left_indent = current_indent + Pt(2)
                        elif col_idx == 4:
                            para.paragraph_format.left_indent = current_indent + Pt(9)
                    else:
                        para.text = line
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        current_indent = para.paragraph_format.left_indent or Pt(0)
                        if col_idx == 2:
                            para.paragraph_format.left_indent = current_indent + Pt(2)
                        elif col_idx == 4:
                            para.paragraph_format.left_indent = current_indent + Pt(9)
                
                if col_idx == 4:
                    set_cell_margins(cell, left=100)
                
                data_index += 1
            else:
                cell = row.cells[col_idx]
                cell.text = ""
    
    print(f"Updated {data_index} data cells from {len(data_rows)} rows")
    
    # Save as temporary DOCX first
    temp_docx = output_path.replace('.pdf', '_temp.docx')
    print(f"\nSaving temporary DOCX to: {temp_docx}")
    doc.save(temp_docx)
    # Convert to PDF
    print(f"Converting to PDF: {output_path}")
    convert_docx_to_pdf(temp_docx, output_path)
    # Remove temporary DOCX
    import os
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
        print(f"Removed temporary file: {temp_docx}")
    print(f"Successfully updated {len(data_rows)} cells and saved as PDF!")


def main():
    if len(sys.argv) != 3:
        print("Usage: python update_template.py <input_image> <template.docx>")
        print("\nExample:")
        print("  python update_template.py input.png template.docx")
        print("\nThe PDF will be saved with today's date and time (e.g., 2026-02-09_03:45 PM.pdf)")
        sys.exit(1)
    
    image_path = sys.argv[1]
    template_path = sys.argv[2]
    
    # Generate output filename with today's date and time (e.g., 2026-02-09_03:45 PM.pdf)
    now = datetime.now()
    date_time = now.strftime("%Y-%m-%d_%I:%M %p")
    output_path = f"{date_time}.pdf"
    
    # Check if files exist
    if not Path(image_path).exists():
        print(f"Error: Image file not found: {image_path}")
        sys.exit(1)
    
    if not Path(template_path).exists():
        print(f"Error: Template file not found: {template_path}")
        sys.exit(1)
    
    # Extract data from image
    data_rows = extract_table_data_from_image(image_path)
    
    if not data_rows:
        print("Warning: No data extracted from image!")
        sys.exit(1)
    
    # Update template and convert to PDF
    update_template_with_data(template_path, output_path, data_rows)


if __name__ == "__main__":
    main()
