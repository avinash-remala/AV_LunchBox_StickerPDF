"""
PDF/Document generation module.
Generates Word documents and PDFs from order data.
"""

import sys
import io
from pathlib import Path
from typing import List, Optional
from copy import deepcopy
import subprocess
import platform

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches, Emu, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsdecls
    from docx.enum.dml import MSO_THEME_COLOR
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", 
                          "python-docx", "--break-system-packages"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Inches, Emu, Cm
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn, nsdecls
    from docx.enum.dml import MSO_THEME_COLOR
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

from ..core.models import Order
from ..core.markers import get_marker_for_box_rice
from ..utils.watermark import ensure_watermark
from ..config.logging_config import get_logger

log = get_logger("core.pdf_generator")


class PDFGenerator:
    """Generates PDF documents from order data."""
    
    def __init__(self, template_path: str, watermark_enabled: bool = True):
        """
        Initialize the PDF generator.
        
        Args:
            template_path: Path to the Word template file
            watermark_enabled: Whether to add a watermark behind cell text
        """
        self.template_path = template_path
        self.watermark_enabled = watermark_enabled
        self._watermark_path: Optional[str] = None
        
        # Pre-generate watermark if enabled
        if self.watermark_enabled:
            self._watermark_path = ensure_watermark()
            if self._watermark_path:
                log.info(f"Watermark ready: {self._watermark_path}")
            else:
                log.info("No logo found — watermark disabled (place logo at assets/logo.png)")
    
    def generate(self, orders: List[Order], output_path: str) -> bool:
        """
        Generate a PDF from orders.
        
        Args:
            orders: List of Order objects
            output_path: Path for the output PDF file
        
        Returns:
            True if successful, False otherwise
        """
        try:
            # First generate the DOCX
            docx_path = str(output_path).replace('.pdf', '.docx')
            self._generate_docx(orders, docx_path)
            
            # Then convert to PDF
            return self._convert_docx_to_pdf(docx_path, output_path)
        
        except Exception as e:
            print(f"Error generating PDF: {e}")
            return False
    
    def _generate_docx(self, orders: List[Order], output_path: str) -> bool:
        """
        Generate a Word document from orders.
        
        Args:
            orders: List of Order objects
            output_path: Path for the output DOCX file
        
        Returns:
            True if successful, False otherwise
        """
        log.info(f"Loading template: {self.template_path}")
        doc = Document(self.template_path)
        table = doc.tables[0]
        data_columns = [0, 2, 4]
        
        # Convert orders to data rows
        data_rows = [
            {
                'name': order.name,
                'address': order.address,
                'box_type': order.box_type,
                'rice_type': order.rice_type,
            }
            for order in orders
        ]
        
        initial_row_count = len(table.rows)
        log.info(f"Template has {initial_row_count} rows initially")
        
        # Add rows if needed
        rows_needed = len(data_rows)
        if rows_needed > initial_row_count:
            additional_rows_needed = rows_needed - initial_row_count
            log.info(f"Need to add {additional_rows_needed} new rows for {len(data_rows)} data items")
            
            if initial_row_count > 0:
                reference_row = table.rows[0]
                tbl = table._element
                
                for _ in range(additional_rows_needed):
                    new_row = deepcopy(reference_row._element)
                    tbl.append(new_row)
        
        # Fill rows with data
        data_index = 0
        for row_idx, row in enumerate(table.rows):
            for col_idx in data_columns:
                if data_index < len(data_rows):
                    data = data_rows[data_index]
                    cell_text = f"{data['name']}\n{data['address']}\n{data['box_type']} - {data['rice_type']}"
                    
                    cell = row.cells[col_idx]
                    cell.text = ""
                    
                    # Track whether we need to insert the watermark on the first paragraph
                    needs_watermark = (self._watermark_path and self.watermark_enabled)
                    
                    for idx, line in enumerate(cell_text.split('\n')):
                        if idx > 0:
                            cell.add_paragraph()
                        para = cell.paragraphs[idx]
                        
                        if idx == 2:
                            # Box/rice line — centered, with marker
                            para.clear()
                            run = para.add_run(line)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Get marker based on box/rice combination
                            marker, font_size_increase = get_marker_for_box_rice(
                                data['box_type'], data['rice_type']
                            )
                            
                            current_font_size = run.font.size
                            if current_font_size is None:
                                current_font_size = Pt(11)
                            elif isinstance(current_font_size, int):
                                current_font_size = Pt(current_font_size / 100)
                            
                            if marker:
                                marker_para = cell.add_paragraph()
                                marker_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                marker_run = marker_para.add_run(marker)
                                marker_run.bold = True
                                marker_run.font.size = current_font_size + Pt(font_size_increase)
                        elif idx == 0:
                            # Name line — bold, centered
                            para.clear()
                            run = para.add_run(line)
                            run.bold = True
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            current_indent = para.paragraph_format.left_indent or Pt(0)
                            if col_idx == 2:
                                para.paragraph_format.left_indent = current_indent + Pt(2)
                            elif col_idx == 4:
                                para.paragraph_format.left_indent = current_indent + Pt(9)
                            # Insert watermark behind text on the first paragraph
                            if needs_watermark:
                                self._insert_watermark_into_cell(doc, cell)
                        else:
                            # Address line — centered
                            para.text = line
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            current_indent = para.paragraph_format.left_indent or Pt(0)
                            if col_idx == 2:
                                para.paragraph_format.left_indent = current_indent + Pt(2)
                            elif col_idx == 4:
                                para.paragraph_format.left_indent = current_indent + Pt(9)
                    
                    if col_idx == 4:
                        self._set_cell_margins(cell, left=100)
                    
                    data_index += 1
                else:
                    cell = row.cells[col_idx]
                    cell.text = ""
        
        # Save DOCX
        doc.save(output_path)
        print(f"DOCX file saved as: {output_path}")
        return True
    
    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """
        Convert DOCX to PDF using available tools.
        
        Args:
            docx_path: Path to the DOCX file
            pdf_path: Path for the output PDF file
        
        Returns:
            True if successful, False otherwise
        """
        system = platform.system()
        
        try:
            if system == "Darwin":  # macOS
                # Try LibreOffice
                result = subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(Path(pdf_path).parent),
                    docx_path
                ], capture_output=True, timeout=30)
                
                if result.returncode == 0:
                    print(f"PDF converted successfully to: {pdf_path}")
                    return True
            
            elif system == "Windows":
                # Try LibreOffice
                result = subprocess.run([
                    "soffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(Path(pdf_path).parent),
                    docx_path
                ], capture_output=True, timeout=30)
                
                if result.returncode == 0:
                    print(f"PDF converted successfully to: {pdf_path}")
                    return True
            
            elif system == "Linux":
                # Try LibreOffice
                result = subprocess.run([
                    "libreoffice", "--headless", "--convert-to", "pdf",
                    "--outdir", str(Path(pdf_path).parent),
                    docx_path
                ], capture_output=True, timeout=30)
                
                if result.returncode == 0:
                    print(f"PDF converted successfully to: {pdf_path}")
                    return True
        
        except Exception as e:
            print(f"Error converting to PDF: {e}")
        
        # Fallback: just keep DOCX
        print(f"Could not convert to PDF. DOCX file available at: {docx_path}")
        return False

    @staticmethod
    def _set_cell_margins(cell, left=None, right=None, top=None, bottom=None):
        """
        Set cell margins (padding) in twentieths of a point (twips).
        1 point = 20 twips.
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')

        margins = {'left': left, 'right': right, 'top': top, 'bottom': bottom}
        for margin_name, value in margins.items():
            if value is not None:
                node = OxmlElement(f'w:{margin_name}')
                node.set(qn('w:w'), str(value))
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

        tcPr.append(tcMar)

    def _insert_watermark_into_cell(self, doc, cell):
        """
        Insert a watermark image into a cell as a behind-text anchored image.
        Sized and positioned to fit centred within the cell (6.47cm × 2.54cm).

        Args:
            doc: The Document object (needed for adding image relationships)
            cell: The table cell to add the watermark to
        """
        if not self._watermark_path or not Path(self._watermark_path).exists():
            return

        para = cell.paragraphs[0]
        part = para.part

        # Use python-docx's built-in image handling: returns (rId, Image)
        rel_id, image = part.get_or_add_image(self._watermark_path)
        img_width_px = image.px_width
        img_height_px = image.px_height

        # Cell is ~6.47cm wide × 2.54cm tall.
        # Fit watermark to ~80% of cell height, then reduce by 20%.
        max_height = Cm(2.0 * 0.8)   # 20% smaller
        max_width  = Cm(5.0 * 0.8)   # 20% smaller

        aspect = img_width_px / img_height_px
        # Height-constrained sizing
        cy = int(max_height)
        cx = int(cy * aspect)
        # If too wide, constrain by width instead
        if cx > int(max_width):
            cx = int(max_width)
            cy = int(cx / aspect)

        # Left-aligned with 3px margin from left border
        # 3px ≈ 2.25pt ≈ 28575 EMU
        h_offset = 28575
        # Top-aligned with small margin (~3px) from top border
        v_offset = 28575

        # Build the anchor XML: behind text, fixed position in cell
        unique_id = id(cell) & 0xFFFFFFFF
        anchor_xml = (
            f'<wp:anchor distT="0" distB="0" distL="0" distR="0" '
            f'simplePos="0" relativeHeight="0" behindDoc="1" locked="1" '
            f'layoutInCell="1" allowOverlap="0" '
            f'{nsdecls("wp", "a", "pic", "r")}>'
            f'  <wp:simplePos x="0" y="0"/>'
            f'  <wp:positionH relativeFrom="column">'
            f'    <wp:posOffset>{h_offset}</wp:posOffset>'
            f'  </wp:positionH>'
            f'  <wp:positionV relativeFrom="paragraph">'
            f'    <wp:posOffset>{v_offset}</wp:posOffset>'
            f'  </wp:positionV>'
            f'  <wp:extent cx="{cx}" cy="{cy}"/>'
            f'  <wp:effectExtent l="0" t="0" r="0" b="0"/>'
            f'  <wp:wrapNone/>'
            f'  <wp:docPr id="{unique_id}" name="Watermark {unique_id}"/>'
            f'  <wp:cNvGraphicFramePr>'
            f'    <a:graphicFrameLocks noChangeAspect="1"/>'
            f'  </wp:cNvGraphicFramePr>'
            f'  <a:graphic>'
            f'    <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            f'      <pic:pic>'
            f'        <pic:nvPicPr>'
            f'          <pic:cNvPr id="{unique_id}" name="watermark.png"/>'
            f'          <pic:cNvPicPr/>'
            f'        </pic:nvPicPr>'
            f'        <pic:blipFill>'
            f'          <a:blip r:embed="{rel_id}"/>'
            f'          <a:stretch>'
            f'            <a:fillRect/>'
            f'          </a:stretch>'
            f'        </pic:blipFill>'
            f'        <pic:spPr>'
            f'          <a:xfrm>'
            f'            <a:off x="0" y="0"/>'
            f'            <a:ext cx="{cx}" cy="{cy}"/>'
            f'          </a:xfrm>'
            f'          <a:prstGeom prst="rect">'
            f'            <a:avLst/>'
            f'          </a:prstGeom>'
            f'        </pic:spPr>'
            f'      </pic:pic>'
            f'    </a:graphicData>'
            f'  </a:graphic>'
            f'</wp:anchor>'
        )

        from lxml import etree
        anchor_element = etree.fromstring(anchor_xml)

        drawing = OxmlElement('w:drawing')
        drawing.append(anchor_element)

        # Append to the first run in this paragraph
        if not para.runs:
            run_element = OxmlElement('w:r')
            para._element.append(run_element)
        else:
            run_element = para.runs[0]._element

        run_element.append(drawing)
