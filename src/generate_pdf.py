#!/usr/bin/env python3
"""
Generate lunch box sticker PDFs and summaries.

Usage:
    python3 src/generate_pdf.py templates/AR_Template.docx --google-sheet <spreadsheet_id>
    python3 src/generate_pdf.py templates/AR_Template.docx --image <image_path>
"""

import sys
import re
import os
import csv
import shutil
import subprocess
from io import StringIO
from copy import deepcopy
from pathlib import Path
from datetime import datetime
from zoneinfo import ZoneInfo

CST = ZoneInfo("America/Chicago")
from collections import Counter

import requests

try:
    from PIL import Image
    import pytesseract
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pillow", "pytesseract", "--break-system-packages"])
    from PIL import Image
    import pytesseract

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Google Sheets
# ---------------------------------------------------------------------------

def get_todays_lunch_orders(spreadsheet_id, sheet_id=0):
    """Fetch today's lunch orders from a public Google Sheet."""
    url = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv&gid={sheet_id}"
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        rows = list(csv.DictReader(StringIO(response.text)))
        print(f"Fetched {len(rows)} rows from Google Sheet")
        if rows:
            print(f"Available columns: {list(rows[0].keys())}")
    except requests.exceptions.RequestException as e:
        print(f"Error fetching Google Sheet: {e}")
        return []

    today = datetime.now(CST)
    today_formats = [today.strftime("%m/%d/%Y"), f"{today.month}/{today.day}/{today.year}"]

    # Find the date column key (handles whitespace-only or renamed headers)
    date_col = 'Date'
    if rows:
        for key in rows[0].keys():
            if key.strip() == 'Date':
                date_col = key
                break
        else:
            # Fall back to first column if no 'Date' header found
            first_key = next(iter(rows[0].keys()), None)
            if first_key is not None:
                date_col = first_key
                print(f"Warning: 'Date' column not found, using first column: {repr(date_col)}")

    orders = []
    found_today = False
    for row in rows:
        date_val = row.get(date_col, '').strip()
        if date_val in today_formats:
            found_today = True
        elif found_today and date_val != '':
            break
        elif not found_today:
            continue

        name = row.get('Full Name', '').strip()
        if name:
            orders.append({
                'name':     name,
                'address':  row.get('Address', '').strip(),
                'box_type': row.get('Type of Food', '').strip(),
                'rice_type':row.get('Type of Rice', '').strip(),
                'comments': row.get('Comments', '').strip(),
            })
            print(f"  Order {len(orders)}: {name} - {orders[-1]['box_type']} - {orders[-1]['rice_type']}")

    print(f"Extracted {len(orders)} total orders from today's data")
    return orders


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

STANDARD_COMBINATIONS = [
    "Veg Comfort Box + Pulav Rice",
    "Non-Veg Comfort Box + Pulav Rice",
    "Veg Comfort Box + White Rice",
    "Non-Veg Comfort Box + White Rice",
]


def save_summary(orders, output_dir, filename):
    """Generate and save a summary text file."""
    if not orders:
        content = "No orders found for this date."
    else:
        box_rice = Counter(f"{o.get('box_type','')} + {o.get('rice_type','')}" for o in orders)
        addr_counts = Counter(o.get('address', '') for o in orders)

        lines = [f"TOTAL BOXES: {len(orders)}\n", "Boxes (count by type)"]
        for combo in STANDARD_COMBINATIONS:
            count = box_rice.get(combo, 0)
            if count > 0:
                lines.append(f"•\t{combo}: {count}")
        for combo, count in sorted(box_rice.items()):
            if combo not in STANDARD_COMBINATIONS:
                lines.append(f"•\t{combo}: {count}")
        lines.append("\nAddresses (total boxes per address)")
        for address in sorted(addr_counts):
            count = addr_counts[address]
            lines.append(f"•\t{address}: {count} {'box' if count == 1 else 'boxes'}")
        content = "\n".join(lines)

    path = Path(output_dir) / filename
    path.parent.mkdir(parents=True, exist_ok=True)
    try:
        path.write_text(content)
        print(f"✓ Summary saved to: {path}")
    except Exception as e:
        print(f"✗ Error saving summary: {e}")


# ---------------------------------------------------------------------------
# PDF helpers
# ---------------------------------------------------------------------------

def get_marker(box_type, rice_type):
    """Return (marker_string, font_size_increase_pt) for a box/rice combo."""
    is_non_veg = "Non-Veg" in box_type
    is_veg = "Veg" in box_type and not is_non_veg
    is_pulav = "Pulav Rice" in rice_type

    if "Special Box" in box_type:
        return ("--- NVSP ---" if is_non_veg else "--- VSP ---"), 2
    if "Comfort Box" in box_type:
        if is_veg and is_pulav:     return "--- VP ---", 2
        if is_non_veg and is_pulav: return "--- NVP ---", 2
        if is_veg:                  return "--- VW ---", 2
        if is_non_veg:              return "--- NVW ---", 2
    return "", 0


def set_cell_margins(cell, left=None, right=None, top=None, bottom=None):
    """Set cell padding in twips (1pt = 20 twips)."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, value in [('left', left), ('right', right), ('top', top), ('bottom', bottom)]:
        if value is not None:
            node = OxmlElement(f'w:{name}')
            node.set(qn('w:w'), str(value))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)


def convert_to_pdf(docx_path, pdf_path):
    """Convert a DOCX file to PDF using LibreOffice."""
    soffice_candidates = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        'soffice', 'libreoffice', '/usr/bin/soffice',
    ]
    for soffice in soffice_candidates:
        try:
            result = subprocess.run(
                [soffice, '--headless', '--convert-to', 'pdf',
                 '--outdir', str(Path(pdf_path).parent), docx_path],
                capture_output=True, timeout=60
            )
            if result.returncode == 0:
                generated = str(Path(pdf_path).parent / Path(docx_path).with_suffix('.pdf').name)
                if Path(generated).exists() and generated != pdf_path:
                    shutil.move(generated, pdf_path)
                if Path(pdf_path).exists():
                    print(f"PDF created successfully with LibreOffice")
                    return
        except Exception:
            continue
    print(f"Could not convert to PDF. DOCX saved at: {docx_path}")


# ---------------------------------------------------------------------------
# OCR (--image mode)
# ---------------------------------------------------------------------------

def extract_from_image(image_path):
    """Extract order data from an image using OCR."""
    print(f"Loading image: {image_path}")
    text = pytesseract.image_to_string(Image.open(image_path))
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    print(f"OCR extracted {len(lines)} lines")

    phone_re   = re.compile(r'\b\d{10}\b')
    address_re = re.compile(r'(2900 Plano Pkwy|3400 W Plano Pkwy)')
    box_re     = re.compile(r'((?:Veg|Non-Veg)\s+Comfort Box|(?:\d+\s+)?(?:Comfort Box|Kabuli Chana Box|Moong Dal Box|Rajma Box))')
    rice_re    = re.compile(r'(Pulav Rice|White Rice)')

    orders, box_types, rice_types = [], [], []

    for line in lines:
        m = address_re.search(line)
        if m:
            address = m.group(1)
            remainder = line.replace(address, '')
            phone = phone_re.search(remainder)
            if phone:
                remainder = remainder.replace(phone.group(), '')
            cleaned = re.sub(r'[»|¥«{}]', ' ', remainder).strip('_- ~{').strip()
            name_parts = re.findall(r'[A-Za-z]+(?:\s+[A-Za-z]+)*', cleaned)
            name = max(name_parts, key=len) if name_parts else ''
            # Split CamelCase
            if name and ' ' not in name and len(name) > 4:
                caps = [i for i, c in enumerate(name) if c.isupper()]
                if len(caps) >= 2:
                    name = ' '.join(name[caps[i]:caps[i+1] if i+1 < len(caps) else None] for i in range(len(caps)))
            parts = name.split()
            while parts and len(parts[0]) <= 2:
                parts.pop(0)
            name = ' '.join(parts)
            if name and len(name) > 3:
                orders.append({'name': name, 'address': address})

        bm = box_re.search(line)
        if bm: box_types.append(bm.group(1))
        rm = rice_re.search(line)
        if rm: rice_types.append(rm.group(1))

    rows = []
    for i in range(max(len(orders), len(box_types), len(rice_types))):
        rows.append({
            'name':      orders[i]['name']    if i < len(orders)     else '',
            'address':   orders[i]['address'] if i < len(orders)     else '',
            'box_type':  box_types[i]         if i < len(box_types)  else '',
            'rice_type': rice_types[i]        if i < len(rice_types) else '',
            'comments':  '',
        })
    print(f"Extracted {len(rows)} rows from image")
    return rows


# ---------------------------------------------------------------------------
# Core PDF generation
# ---------------------------------------------------------------------------

def generate_pdf(template_path, output_path, data_rows):
    """Fill the Word template with order data and convert to PDF."""
    doc = Document(template_path)
    table = doc.tables[0]
    data_columns = [0, 2, 4]

    # Expand table rows if needed
    initial_rows = len(table.rows)
    if len(data_rows) > initial_rows:
        ref = table.rows[0]._element
        for _ in range(len(data_rows) - initial_rows):
            table._element.append(deepcopy(ref))
    print(f"Template has {initial_rows} rows; filling {len(data_rows)} stickers")

    data_index = 0
    for row in table.rows:
        for col_idx in data_columns:
            cell = row.cells[col_idx]
            if data_index >= len(data_rows):
                cell.text = ""
                continue

            data = data_rows[data_index]
            cell.text = ""

            # Line 0 — Name (bold, centered)
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(data['name'])
            run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            indent = para.paragraph_format.left_indent or Pt(0)
            if col_idx == 2: para.paragraph_format.left_indent = indent + Pt(2)
            elif col_idx == 4: para.paragraph_format.left_indent = indent + Pt(9)

            # Line 1 — Address (centered)
            para = cell.add_paragraph()
            para.text = data['address']
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            indent = para.paragraph_format.left_indent or Pt(0)
            if col_idx == 2: para.paragraph_format.left_indent = indent + Pt(2)
            elif col_idx == 4: para.paragraph_format.left_indent = indent + Pt(9)

            # Line 2 — Box type - Rice type (centered)
            para = cell.add_paragraph()
            para.clear()
            run = para.add_run(f"{data['box_type']} - {data['rice_type']}")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font_size = Pt(11)

            # Comments (italic, 2pt smaller) — only if present
            comments = data.get('comments', '').strip()
            if comments:
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(comments)
                run.italic = True
                run.font.size = font_size - Pt(2)

            # Marker (bold, 2pt larger)
            marker, size_inc = get_marker(data['box_type'], data['rice_type'])
            if marker:
                para = cell.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(marker)
                run.bold = True
                run.font.size = font_size + Pt(size_inc)

            if col_idx == 4:
                set_cell_margins(cell, left=100)

            data_index += 1

    print(f"Filled {data_index} sticker cells")

    temp_docx = output_path.replace('.pdf', '_temp.docx')
    doc.save(temp_docx)
    print(f"Saved DOCX: {temp_docx}")
    convert_to_pdf(temp_docx, output_path)
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
    print(f"✓ PDF saved: {output_path}")


def cleanup_old_exports(exports_dir="exports"):
    """Delete old export folders, keeping only today's."""
    exports = Path(exports_dir)
    if not exports.exists():
        return
    today = datetime.now(CST).strftime("%Y-%m-%d")
    for item in exports.iterdir():
        if item.is_dir() and item.name == today:
            continue
        try:
            shutil.rmtree(item) if item.is_dir() else item.unlink()
        except Exception as e:
            print(f"✗ Could not delete {item.name}: {e}")


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    if len(sys.argv) < 2:
        print("Usage: python3 generate_pdf.py <template.docx> --google-sheet <id>")
        print("       python3 generate_pdf.py <template.docx> --image <path>")
        sys.exit(1)

    template_path = sys.argv[1]
    if not Path(template_path).exists():
        print(f"Error: Template not found: {template_path}")
        sys.exit(1)

    if '--google-sheet' in sys.argv:
        idx = sys.argv.index('--google-sheet')
        spreadsheet_id = sys.argv[idx + 1]
        print(f"\nReading data from Google Sheet: {spreadsheet_id}")
        # Wipe all exports before a fresh Google Sheets run
        if Path("exports").exists():
            shutil.rmtree("exports")
        data_rows = get_todays_lunch_orders(spreadsheet_id)

    elif '--image' in sys.argv:
        idx = sys.argv.index('--image')
        image_path = sys.argv[idx + 1]
        if not Path(image_path).exists():
            print(f"Error: Image not found: {image_path}")
            sys.exit(1)
        data_rows = extract_from_image(image_path)

    else:
        print("Error: Provide --google-sheet <id> or --image <path>")
        sys.exit(1)

    if not data_rows:
        print("No data found. Exiting.")
        sys.exit(1)

    now = datetime.now(CST)
    timestamp = now.strftime("%Y-%m-%d_%I-%M%p")
    output_dir = Path("exports") / now.strftime("%Y-%m-%d")
    output_dir.mkdir(parents=True, exist_ok=True)

    cleanup_old_exports("exports")
    generate_pdf(template_path, str(output_dir / f"{timestamp}.pdf"), data_rows)
    save_summary(data_rows, str(output_dir), f"{timestamp}.txt")


if __name__ == "__main__":
    main()
